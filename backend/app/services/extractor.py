"""Document text extraction service.

Phase 4.2 / 4.3 — Extracts raw text from PDF (PyMuPDF) and DOCX (python-docx)
files. All processing is in-memory; no files are written to disk.

Phase 4.4 — Context window chunking for documents exceeding ~8k tokens.
"""

import io
import logging

import docx
import fitz  # PyMuPDF

from ..models.response import DocumentAnalysis

logger = logging.getLogger(__name__)


def extract_pdf_text(file_bytes: bytes) -> DocumentAnalysis:
    """Extract text from a PDF file byte stream.

    Uses PyMuPDF (fitz) for page-by-page text extraction.
    Inserts page markers between pages to preserve document structure.

    Args:
        file_bytes: Raw PDF file contents.

    Returns:
        DocumentAnalysis with raw_text, char_count, page_count, file_type.

    Raises:
        ValueError: If no text could be extracted (scanned/image PDF).
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages.append(text)

    doc.close()

    if not pages:
        raise ValueError(
            "We couldn't read any text from this file. "
            "Please check it's a text-based PDF or DOCX, not a scanned image."
        )

    # Join with page markers for structure preservation
    raw_text = ""
    for i, page_text in enumerate(pages):
        if i > 0:
            raw_text += f"\n--- PAGE {i + 1} ---\n"
        raw_text += page_text

    logger.info(
        "PDF extraction complete: %d pages, %d characters",
        len(pages),
        len(raw_text),
    )

    return DocumentAnalysis(
        raw_text=raw_text,
        char_count=len(raw_text),
        page_count=len(pages),
        file_type="pdf",
    )


def extract_docx_text(file_bytes: bytes) -> DocumentAnalysis:
    """Extract text from a DOCX file byte stream.

    Uses python-docx to iterate paragraphs and table cells in document
    order, preserving the full document structure (TRD §3.3).

    Args:
        file_bytes: Raw DOCX file contents.

    Returns:
        DocumentAnalysis with raw_text, char_count, page_count=None, file_type.

    Raises:
        ValueError: If no text could be extracted.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cell text in document order
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    if not parts:
        raise ValueError(
            "We couldn't read any text from this file. "
            "Please check it's a text-based PDF or DOCX, not a scanned image."
        )

    raw_text = "\n".join(parts)

    logger.info(
        "DOCX extraction complete: %d text segments, %d characters",
        len(parts),
        len(raw_text),
    )

    return DocumentAnalysis(
        raw_text=raw_text,
        char_count=len(raw_text),
        page_count=None,
        file_type="docx",
    )


def extract_text(file_bytes: bytes, file_type: str) -> DocumentAnalysis:
    """Route extraction to the correct handler based on file type.

    Args:
        file_bytes: Raw file contents.
        file_type: Either 'pdf' or 'docx'.

    Returns:
        DocumentAnalysis result.
    """
    if file_type == "pdf":
        return extract_pdf_text(file_bytes)
    return extract_docx_text(file_bytes)


def chunk_text(raw_text: str, max_tokens: int = 5000, overlap_tokens: int = 500) -> list[str]:
    """Split text into overlapping chunks for context window management.

    Phase 4.4 — For documents exceeding ~8k token context window.
    Uses conservative char/4 token approximation.

    Args:
        raw_text: Full document text.
        max_tokens: Maximum tokens per chunk (default 5000).
        overlap_tokens: Token overlap between chunks (default 500).

    Returns:
        List of text chunks, each within the token limit.
    """
    # Conservative token estimate: 1 token ≈ 4 characters
    max_chars = max_tokens * 4
    overlap_chars = overlap_tokens * 4

    if len(raw_text) <= max_chars:
        return [raw_text]

    chunks: list[str] = []
    start = 0

    while start < len(raw_text):
        end = start + max_chars

        # Try to break at a sentence boundary
        if end < len(raw_text):
            # Look back from end for a period, newline, or other boundary
            boundary = raw_text.rfind(". ", start + max_chars // 2, end)
            if boundary == -1:
                boundary = raw_text.rfind("\n", start + max_chars // 2, end)
            if boundary != -1:
                end = boundary + 1

        chunks.append(raw_text[start:end].strip())
        start = end - overlap_chars

    logger.info("Document chunked into %d segments", len(chunks))
    return chunks
